#!/usr/bin/env python3
"""Build the SPARTA teaching chapter as a publication-ready Word document."""

from __future__ import annotations

import csv
import json
import math
import os
from pathlib import Path

os.environ.setdefault("MPLCONFIGDIR", str(Path(__file__).resolve().parent / ".mplconfig"))

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
OUTPUT = DOCS / "Chapter_SPARTA_Mohammadzadeh_Cavity.docx"

BLUE = "2E74B5"
DARK_BLUE = "17365D"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F8FC"
GRAY = "F3F5F7"
DARK_GRAY = "4D4D4D"
ORANGE = "ED7D31"
GREEN = "548235"
RED = "C00000"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B7C9DC", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_widths(table, widths_twips: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_twips):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text: str, url: str, color=BLUE) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PAGE ")
    run.font.size = Pt(8)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, value, end):
        run._r.append(node)


def add_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_top_rule(paragraph, color=BLUE, size="18") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    doc.settings.odd_and_even_pages_header_footer = True
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("202020")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Title", 29, DARK_BLUE, 0, 14),
        ("Subtitle", 14, DARK_GRAY, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        if name == "Heading 1":
            style.paragraph_format.page_break_before = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code_style = doc.styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor.from_string("243447")
    code_style.paragraph_format.left_indent = Inches(0.12)
    code_style.paragraph_format.right_indent = Inches(0.12)
    code_style.paragraph_format.space_before = Pt(5)
    code_style.paragraph_format.space_after = Pt(7)
    code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    code_style.paragraph_format.keep_together = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    header = section.header.paragraphs[0]
    header.text = "PRACTICAL DSMC  •  SPARTA CAVITY BENCHMARK"
    header.style = doc.styles["Normal"]
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    add_top_rule(header, BLUE, "8")
    even_header = section.even_page_header.paragraphs[0]
    even_header.text = "PRACTICAL DSMC  •  SPARTA CAVITY BENCHMARK"
    even_header.style = doc.styles["Normal"]
    even_header.runs[0].font.size = Pt(8)
    even_header.runs[0].font.bold = True
    even_header.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    add_top_rule(even_header, BLUE, "8")
    add_page_field(section.footer.paragraphs[0])


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), GRAY)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:color"), BLUE)
    left.set(qn("w:space"), "5")
    borders.append(left)
    p_pr.append(borders)
    run = p.add_run(text.rstrip())
    run.font.name = "Consolas"


def add_callout(doc: Document, title: str, body: str, color=LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))
    set_table_widths(table, [9360])
    set_table_borders(table, color=BLUE, size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))
    set_table_widths(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cell = cells[idx]
            if row_idx % 2:
                set_cell_shading(cell, "F8FAFC")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.size = Pt(9.5)
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            set_cell_margins(cell)


def add_bullet(doc: Document, text: str, numbered=False) -> None:
    p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    p.add_run(" ")
    p.add_run(text)


def add_figure(doc: Document, path: Path, caption: str, alt: str, width=6.25) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    add_alt_text(shape, caption.split(".")[0], alt)
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def page_break(doc: Document) -> None:
    # Heading 1 carries page_break_before. A standalone break paragraph can be
    # pushed to the next page and create a blank sheet when the prior page is full.
    return None


def format_inline_code(doc: Document) -> None:
    """Replace Markdown-style backticks with compact inline code formatting."""

    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        if paragraph.style and paragraph.style.name == "Code Block":
            continue
        if paragraph._p.xpath(".//w:hyperlink"):
            continue
        text = paragraph.text
        if "`" not in text or text.count("`") % 2:
            continue
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        for index, piece in enumerate(text.split("`")):
            if not piece:
                continue
            run = paragraph.add_run(piece)
            if index % 2:
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor.from_string("1F4D78")


def generate_geometry_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    ax.add_patch(plt.Rectangle((0, 0), 1, 1, fill=False, lw=3, color="#17365D"))
    ax.arrow(0.18, 1.08, 0.58, 0, head_width=0.055, head_length=0.08,
             color="#ED7D31", lw=3, length_includes_head=True)
    ax.text(0.48, 1.16, r"moving lid: $U_w=100$ m/s, $T_w=300$ K",
            ha="center", va="bottom", fontsize=12, color="#9E480E")
    ax.text(0.5, 0.5, "Argon DSMC\nKn = 0.1 (primary case)", ha="center", va="center",
            fontsize=14, color="#17365D", weight="bold")
    ax.text(-0.06, 0.5, "diffuse wall", rotation=90, ha="center", va="center", fontsize=10)
    ax.text(1.06, 0.5, "diffuse wall", rotation=-90, ha="center", va="center", fontsize=10)
    ax.text(0.5, -0.08, "diffuse wall, 300 K", ha="center", va="top", fontsize=10)
    ax.annotate("sample top-cell centers", xy=(0.82, 0.98), xytext=(0.64, 0.68),
                arrowprops=dict(arrowstyle="->", color="#548235", lw=1.8),
                fontsize=10, color="#548235")
    ax.set_xlim(-0.18, 1.18)
    ax.set_ylim(-0.18, 1.32)
    ax.set_aspect("equal")
    ax.axis("off")
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def generate_student_figure(path: Path) -> dict[str, float]:
    profile = ROOT / "runs" / "student_kn01" / "lid_profile.csv"
    metrics_path = ROOT / "runs" / "student_kn01" / "validation_metrics.json"
    reference = ROOT / "reference" / "mohammadzadeh_2012_lid_profiles.csv"
    if not profile.exists() or not metrics_path.exists():
        raise FileNotFoundError("Run the documented student case before building the chapter")
    data = np.genfromtxt(profile, delimiter=",", names=True)
    with reference.open(newline="", encoding="utf-8") as handle:
        ref_rows = [row for row in csv.DictReader(handle) if math.isclose(float(row["kn"]), 0.1)]
    xr = np.array([float(row["x_over_L"]) for row in ref_rows])
    sr = np.array([float(row["macro_slip_over_Uwall"]) for row in ref_rows])
    tr = np.array([float(row["macro_T_K"]) for row in ref_rows])
    metrics = json.loads(metrics_path.read_text(encoding="utf-8"))
    fig, axes = plt.subplots(1, 2, figsize=(9.2, 3.6))
    axes[0].plot(data["x_over_L"], data["macro_slip_over_Uwall"], color="#2E74B5", lw=1.3,
                 label="SPARTA student")
    axes[0].scatter(xr, sr, color="#ED7D31", edgecolor="white", s=42, zorder=3, label="PRE digitization")
    axes[0].set(xlabel="x/L", ylabel=r"$(U_w-u)/U_w$", title="Lid-adjacent velocity slip")
    axes[0].grid(alpha=0.22)
    axes[0].legend(frameon=False, fontsize=8)
    axes[1].plot(data["x_over_L"], data["macro_T_K"], color="#2E74B5", lw=1.3,
                 label="SPARTA student")
    axes[1].scatter(xr, tr, color="#ED7D31", edgecolor="white", s=42, zorder=3, label="PRE digitization")
    axes[1].set(xlabel="x/L", ylabel="T (K)", title="Lid-adjacent temperature")
    axes[1].grid(alpha=0.22)
    axes[1].legend(frameon=False, fontsize=8)
    fig.suptitle("Learning-resolution result — not a production validation", color="#C00000", weight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return metrics


def configure_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Hands-On SPARTA: Reproducing the Mohammadzadeh Lid-Driven Cavity"
    props.subject = "Practical DSMC chapter for the second edition"
    props.author = "Ehsan Roohi"
    props.keywords = "SPARTA, DSMC, rarefied gas, lid-driven cavity, validation, Codex"
    props.comments = "Generated from the reproducible GitHub teaching case."


def build_document() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    geometry = ASSETS / "cavity_geometry.png"
    pilot = ASSETS / "student_pilot_comparison.png"
    generate_geometry_figure(geometry)
    metrics = generate_student_figure(pilot)

    doc = Document()
    style_document(doc)
    configure_properties(doc)

    # Editorial cover
    p = doc.add_paragraph("PRACTICAL CHAPTER  •  SECOND EDITION 2027")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(56)
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(ORANGE)
    add_top_rule(p, ORANGE, "22")
    title = doc.add_paragraph(style="Title")
    title.add_run("Hands-On SPARTA").bold = True
    title.add_run("\nReproducing the Mohammadzadeh\nLid-Driven Cavity")
    subtitle = doc.add_paragraph(
        "A reproducible Linux, MPI, validation, and Codex workflow for students",
        style="Subtitle",
    )
    subtitle.paragraph_format.space_after = Pt(30)
    add_callout(
        doc,
        "Evidence status",
        "The serial smoke workflow and the classroom-resolution case have been executed with official SPARTA source. "
        "The publication-resolution, repeated-seed validation remains explicitly pending.",
        PALE_BLUE,
    )
    doc.add_paragraph("Ehsan Roohi", style="Heading 2")
    p = doc.add_paragraph("Companion repository  •  Ehsan-Roohi/DSMC_Python")
    p.runs[0].font.color.rgb = RGBColor.from_string(DARK_GRAY)
    p = doc.add_paragraph("Draft prepared 3 August 2026")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor.from_string(DARK_GRAY)

    page_break(doc)
    doc.add_heading("How to use this chapter", level=1)
    add_callout(
        doc,
        "Editorial placement",
        "Place this as a standalone practical chapter immediately after the book's theoretical cavity chapter. "
        "Do not bury it at the end of the data-driven chapter: the reader should first learn the physics and then reproduce it in SPARTA before moving to AI-assisted analysis.",
    )
    doc.add_heading("Learning outcomes", level=2)
    for item in (
        "Translate a literature DSMC benchmark into a traceable SPARTA input deck.",
        "Build official SPARTA source locally on Linux in serial and MPI modes.",
        "Distinguish a syntax smoke test, a classroom run, and a publication validation.",
        "Extract lid-adjacent velocity slip and temperature and compare them with digitized reference data.",
        "Use Codex to inspect, run, and explain the workflow without silently changing the benchmark.",
    ):
        add_bullet(doc, item)
    doc.add_heading("Contents", level=2)
    add_table(
        doc,
        ["Part", "Purpose"],
        [
            ["1–2", "Benchmark physics and the SPARTA mapping"],
            ["3–5", "Linux installation, case anatomy, and execution"],
            ["6–8", "Post-processing, validation, and collision-model studies"],
            ["9–11", "Codex workflow, exercises, troubleshooting, and references"],
        ],
        [1500, 7860],
    )
    doc.add_heading("Evidence vocabulary used throughout", level=2)
    add_table(
        doc,
        ["Label", "What it proves", "What it does not prove"],
        [
            ["Smoke", "Deck parses; particles, collisions, walls, averaging, dump, and parser work.", "Agreement with the paper."],
            ["Student", "A usable learning profile and a first quantitative diagnostic.", "Grid/time/PPC independence."],
            ["Production", "A candidate high-resolution result.", "Validation until repeats and uncertainty are reported."],
        ],
        [1450, 3900, 4010],
    )

    page_break(doc)
    doc.add_heading("1. Benchmark definition", level=1)
    add_figure(
        doc,
        geometry,
        "Figure 1. Dimensional realization of the square lid-driven cavity used in this chapter.",
        "Square cavity with a right-moving top lid at 100 metres per second, all walls at 300 kelvin, and diffuse reflection.",
    )
    doc.add_paragraph(
        "Mohammadzadeh et al. studied monatomic argon in a square micro- or nanocavity using DSMC. "
        "All walls are isothermal at 300 K; the top wall translates at 100 m/s. Molecular reflection is fully diffuse with full thermal accommodation. "
        "The paper reports Kn = 0.005, 0.05, and 0.1. This implementation starts with Kn = 0.1 because the rarefaction signal is large and the production grid is comfortably finer than the mean free path."
    )
    add_table(
        doc,
        ["Parameter", "Value", "Traceability"],
        [
            ["Gas / molecular mass", "Ar / 6.63 × 10⁻²⁶ kg", "PRE paper"],
            ["Reference diameter", "4.17 × 10⁻¹⁰ m", "PRE paper"],
            ["Wall temperature", "300 K", "PRE paper"],
            ["Lid speed", "100 m/s", "PRE paper"],
            ["Wall scattering", "Diffuse; accommodation = 1", "PRE paper"],
            ["Collision model", "VHS, ω = 0.81, α = 1", "Paper + repository model file"],
            ["Primary resolution", "200 × 200; 32 particles/cell", "Paper-selected grid and population"],
        ],
        [2500, 2700, 4160],
    )
    doc.add_heading("1.1 Similarity-preserving dimensionalization", level=2)
    doc.add_paragraph(
        "The paper specifies the flow through Kn = λ/L but does not state one mandatory absolute cavity length. "
        "We choose L = 1 μm. This is a dimensional realization, not a new physical claim. Once Kn, wall speed, wall temperature, gas model, and geometry are preserved, L determines λ and the required number density."
    )
    add_code(doc, "lambda = Kn L\nn = 1 / (sqrt(2) pi d_ref^2 lambda)\nfnum = n L^2 / N_sim       # SPARTA 2-D unit-depth convention")
    add_callout(
        doc,
        "Important benchmark nuance",
        "The paper states Δx ≈ 0.1λ and also selects a 200 × 200 grid. Across three Kn values those two statements cannot both hold with one fixed grid. "
        "This package records Δx/λ in every metadata file. For Kn = 0.1, the 200 × 200 production deck gives Δx/λ = 0.05.",
        "FFF2CC",
    )

    page_break(doc)
    doc.add_heading("2. From DSMC physics to SPARTA commands", level=1)
    doc.add_paragraph(
        "SPARTA advances simulator particles through free flight, boundary interaction, and stochastic collisions. "
        "The mapping below makes each literature assumption auditable in the input deck."
    )
    add_table(
        doc,
        ["Physical choice", "SPARTA realization", "Audit question"],
        [
            ["Two-dimensional square", "dimension 2; boundary s s p", "Are x/y walls surfaces and z periodic?"],
            ["Uniform collision cells", "create_grid Nx Nx 1", "Is Δx/λ reported?"],
            ["Argon at 300 K", "species + mixture + global nrho/temp", "Do mass, density, and fnum agree?"],
            ["Fully diffuse fixed walls", "surf_collide fixed diffuse 300 1", "Is accommodation exactly one?"],
            ["Moving diffuse lid", "surf_collide lid … translate 100 0 0", "Is only yhi assigned to lid?"],
            ["VHS/NTC collisions", "collide vss gas argon.vss; α = 1", "Does the data file represent VHS?"],
            ["Macroscopic sampling", "compute grid + fix ave/grid", "Is warm-up excluded from averaging?"],
        ],
        [2600, 3550, 3210],
    )
    doc.add_heading("2.1 Time-step rule", level=2)
    doc.add_paragraph(
        "The generator uses a conservative cell-crossing scale and records both Δt and Δt/τ. For the production preset, Δt = 8.26 × 10⁻¹³ s. "
        "Never treat one time step as universal: repeat the result at a smaller Δt before publication."
    )
    add_code(doc, "c_mp = sqrt(2 k_B T / m)\ndt = 0.25 (L/Nx) / (4 c_mp + U_lid)\ntau = lambda / c_mp")
    doc.add_heading("2.2 Particle weight in a 2-D calculation", level=2)
    doc.add_paragraph(
        "SPARTA's 2-D calculation is interpreted with unit out-of-plane depth. Therefore `fnum` scales with nL² rather than nL³. "
        "The generator targets an initial particles-per-cell value; stochastic rounding may produce one fewer or one more particle globally."
    )

    page_break(doc)
    doc.add_heading("3. Obtain and build the case on Linux", level=1)
    doc.add_heading("3.1 Clone the teaching branch", level=2)
    add_code(doc, "git clone --branch agent/validated-dsmc-cavity --single-branch \\\n  https://github.com/Ehsan-Roohi/DSMC_Python.git\ncd DSMC_Python/sparta_cavity_mohammadzadeh")
    doc.add_heading("3.2 Install prerequisites", level=2)
    add_code(doc, "sudo apt update\nsudo apt install -y git build-essential openmpi-bin libopenmpi-dev \\\n  python3 python3-venv")
    doc.add_paragraph(
        "On a managed cluster, use the site compiler and MPI modules instead of `sudo`. The package installer itself never requests administrative privileges."
    )
    doc.add_heading("3.3 Build pinned official SPARTA source", level=2)
    add_code(doc, "bash scripts/install_sparta_linux.sh serial\n# Optional MPI executable\nbash scripts/install_sparta_linux.sh mpi")
    doc.add_paragraph(
        "The installer clones the official `sparta/sparta` repository, checks out commit `912c9e163c38ea5c3562d039e65215f6e2a4f3f8`, and invokes SPARTA's make target. "
        "Pinning separates a reproducible book case from future upstream changes."
    )
    add_callout(
        doc,
        "One-command bootstrap",
        "From an empty directory, the documented bootstrap clones this branch, builds serial SPARTA, runs the unit test, and executes the smoke case. "
        "It refuses to overwrite an existing directory.",
    )
    add_code(doc, "bash <(curl -fsSL https://raw.githubusercontent.com/Ehsan-Roohi/DSMC_Python/agent/validated-dsmc-cavity/sparta_cavity_mohammadzadeh/scripts/bootstrap_linux.sh)")
    doc.add_heading("3.4 Repository map", level=2)
    add_code(doc, "sparta_cavity_mohammadzadeh/\n├── cases/                 # committed generated decks\n├── data/                  # argon species and collision data\n├── reference/             # digitized PRE profiles\n├── scripts/               # generate, run, install, post-process\n├── hpc/                   # Unity CPU/MPI template\n├── tests/                 # regression checks\n├── AGENTS.md              # constraints for Codex\n└── VALIDATION_STATUS.md   # current scientific evidence")

    page_break(doc)
    doc.add_heading("4. Read the SPARTA input deck", level=1)
    doc.add_paragraph(
        "The deck is generated from metadata; do not hand-edit derived values. The first block defines units, stochastic seed, geometry, computational grid, and global gas state."
    )
    add_code(doc, "units                si\nseed                 20260803\ndimension            2\nboundary             s s p\n\ncreate_box           0.0 1.0e-6 0.0 1.0e-6 -0.5 0.5\ncreate_grid          200 200 1\n\nglobal               nrho 1.2943836530e25 fnum 1.011237229e7 temp 300\nspecies              argon.species Ar\nmixture              gas Ar nrho 1.2943836530e25 temp 300 vstream 0 0 0")
    doc.add_heading("4.1 Why `boundary s s p`?", level=2)
    doc.add_paragraph(
        "For a 2-D problem, x and y are physical surface boundaries (`s`) while z is periodic (`p`). The z extent in `create_box` is bookkeeping for the 2-D unit-depth formulation."
    )
    doc.add_heading("4.2 Wall and lid models", level=2)
    add_code(doc, "surf_collide fixed diffuse 300 1.0\nsurf_collide lid   diffuse 300 1.0 translate 100 0 0\nbound_modify xlo xhi ylo collide fixed\nbound_modify yhi collide lid")
    doc.add_paragraph(
        "The `translate` keyword supplies the moving wall velocity to the diffuse reflection model. Assignment is explicit: three stationary faces use `fixed`, and only the top face (`yhi`) uses `lid`."
    )
    add_callout(
        doc,
        "Common error",
        "A moving incoming gas stream is not the same as a moving wall. Keep the initial mixture stream at zero and translate the lid surface model.",
        "FCE4D6",
    )

    page_break(doc)
    doc.add_heading("5. Collisions, sampling, and output", level=1)
    doc.add_heading("5.1 VHS collisions", level=2)
    add_code(doc, "collide              vss gas argon.vss\ncreate_particles     gas n 0\ntimestep             8.2567881869e-13")
    doc.add_paragraph(
        "SPARTA's `vss` collision style reads one line per species. Setting the scattering parameter α = 1 gives the VHS angular law. "
        "The `n 0` option lets the specified number density and `fnum` determine the initial simulator population."
    )
    doc.add_heading("5.2 Separate warm-up from sampling", level=2)
    add_code(doc, "run                  14000\n\nreset_timestep       0\ncompute              flow grid all gas nrho u v w temp\nfix                  flowavg ave/grid all 10 1 10 c_flow[*] ave running\ndump                 fields grid all 26000 grid.final.* id xc yc f_flowavg[*]\ndump_modify          fields pad 8\nrun                  26000")
    doc.add_paragraph(
        "No production sample is accumulated during warm-up. `fix ave/grid` forms a running mean during the sampling phase, and the final dump carries cell centers plus averaged density, velocity components, and temperature."
    )
    doc.add_heading("5.3 Presets", level=2)
    add_table(
        doc,
        ["Preset", "Grid", "Particles/cell", "Warm-up + sample", "Δx/λ", "Claim allowed"],
        [
            ["smoke", "12²", "4", "20 + 30", "0.833", "Syntax/data path"],
            ["student", "50²", "12", "4,000 + 8,000", "0.200", "Learning profile"],
            ["production", "200²", "32", "14,000 + 26,000", "0.050", "Validation candidate"],
        ],
        [1250, 1050, 1500, 1900, 1050, 2610],
    )
    doc.add_heading("5.4 Run commands", level=2)
    add_code(doc, "python3 -m unittest discover -s tests -v\nbash scripts/run_case.sh smoke serial\nbash scripts/run_case.sh student serial\n\n# MPI examples\nMPI_RANKS=8  bash scripts/run_case.sh student mpi\nMPI_RANKS=16 bash scripts/run_case.sh production mpi")

    page_break(doc)
    doc.add_heading("6. Read the outputs and compare with the paper", level=1)
    add_figure(
        doc,
        pilot,
        "Figure 2. Executed student-resolution pilot at Kn = 0.1 compared with digitized PRE lid profiles. The red heading is intentional: this is not a production validation.",
        "Two plots compare a student SPARTA run with digitized reference data: velocity-slip agreement is promising, while temperature remains noisy and biased.",
    )
    doc.add_paragraph(
        f"The executed 50 × 50, 12-particles-per-cell pilot produced slip RMSE = {metrics['slip_rmse']:.4f} and temperature RMSE = {metrics['temperature_rmse_K']:.2f} K. "
        "The slip regression gate passed; the temperature gate did not. This is useful diagnostic evidence: the case is wired correctly, but the learning resolution and sample window are insufficient for a publication claim."
    )
    doc.add_heading("6.1 Output contract", level=2)
    add_table(
        doc,
        ["File", "Meaning", "Reader check"],
        [
            ["case_metadata.json", "Inputs plus derived λ, n, fnum, Δt, Δx/λ, and evidence label", "Can the case be regenerated exactly?"],
            ["log.cavity", "SPARTA performance and event counts", "Are collisions and wall hits nonzero?"],
            ["grid.final.*", "Final running-average grid fields", "Are expected columns present?"],
            ["lid_profile.csv", "Top-cell-center slip and temperature", "Is x/L monotonic and complete?"],
            ["validation_metrics.json", "RMSE gates and pass/fail status", "Is a learning run being overclaimed?"],
        ],
        [2200, 4100, 3060],
    )
    doc.add_heading("6.2 Macro sampling versus direct wall sampling", level=2)
    doc.add_paragraph(
        "This package compares macroscopic values in the centers of the cells adjacent to the lid. The PRE paper discusses meaningful differences between direct wall sampling and macroscopic cell sampling near nonequilibrium corners. "
        "Do not mix those two observables in one validation curve. A future extension should add direct incident/reflected molecular sampling as a separate dataset."
    )

    page_break(doc)
    doc.add_heading("7. Publication-validation protocol", level=1)
    doc.add_paragraph(
        "A visually plausible curve is not validation. Use the following protocol and retain every metadata/metrics file."
    )
    steps = (
        "Baseline: run the 200 × 200, 32-particles-per-cell deck at Kn = 0.1.",
        "Repeat at least three statistically independent seeds; report means and 95% confidence intervals.",
        "Grid study: compare at least 100², 200², and 400² while controlling particles per cell and sample quality.",
        "Particle study: increase particles/cell at fixed grid and time step.",
        "Time-step study: reduce Δt; quantify changes rather than quoting a rule only.",
        "Sampling study: increase warm-up and averaging windows; demonstrate stationary means.",
        "Compare raw top-cell-center profiles with the digitized paper data.",
        "Apply the paper's five-neighbor filter only as a second, labelled curve; never discard the raw data.",
        "Cross-check with the repository's independent Python NTC-PreScan solver.",
    )
    for step in steps:
        add_bullet(doc, step, numbered=True)
    doc.add_heading("7.1 Metrics", level=2)
    add_code(doc, "RMSE_q = sqrt[(1/N) sum_i (q_SPARTA(x_i) - q_PRE(x_i))^2]\n\nCurrent regression targets:\n  velocity-slip RMSE <= 0.08\n  temperature RMSE   <= 2 K")
    doc.add_paragraph(
        "Reference coordinates are generally different from grid-cell centers; the post-processor interpolates the SPARTA profile to reference x/L values. "
        "Digitization uncertainty is recorded in the reference CSV and should be propagated in the final book results."
    )
    add_callout(
        doc,
        "Decision rule",
        "Passing the two RMSE gates is necessary for the automated regression, but not sufficient for scientific validation. "
        "Convergence and uncertainty evidence remain mandatory.",
        "FFF2CC",
    )
    doc.add_heading("7.2 Unity cluster", level=2)
    add_code(doc, "cd DSMC_Python/sparta_cavity_mohammadzadeh\nmkdir -p logs\nsbatch hpc/unity_sparta_production.slurm")
    p = doc.add_paragraph(
        "Unity template: 16 CPU/MPI tasks; verify local modules. GPU SPARTA needs a separate Kokkos/CUDA build."
    )
    p.style = doc.styles["Caption"]

    page_break(doc)
    doc.add_heading("8. Collision-model experiments", level=1)
    doc.add_paragraph(
        "The published benchmark baseline is VHS. SPARTA names the general command `collide vss`; the species-pair data determine whether the angular law behaves as VHS or VSS."
    )
    add_table(
        doc,
        ["Experiment", "argon.vss α", "Purpose", "Validation status"],
        [
            ["VHS baseline", "1.0", "Reproduce the paper model", "Eligible after convergence study"],
            ["VSS sensitivity", "1.4", "Test angular-scattering sensitivity", "Not the paper baseline"],
            ["Collisionless", "collide none", "Free-molecular code check only", "Not a cavity benchmark substitute"],
        ],
        [1800, 1550, 3250, 2760],
    )
    doc.add_heading("8.1 Make a model comparison reproducible", level=2)
    for item in (
        "Copy the baseline case into a new named run directory; never overwrite baseline output.",
        "Change only one collision-model parameter at a time.",
        "Record the model identifier and source data in metadata.",
        "Use the same seed set, grid, time step, population, warm-up, and sample window.",
        "Compare profiles with uncertainty bands, not only one-seed lines.",
    ):
        add_bullet(doc, item)
    add_code(doc, "# Baseline argon.vss\n# species  diameter(m)  omega  Tref(K)  alpha\nAr         4.17e-10     0.81   273.0    1.0\n\n# Sensitivity-only VSS variant\nAr         4.17e-10     0.81   273.0    1.4")
    add_callout(
        doc,
        "Do not move the goalposts",
        "If the VSS sensitivity happens to fit one observable better, it must still be labelled a different physical model. "
        "It does not retroactively become the Mohammadzadeh VHS validation.",
        "FCE4D6",
    )

    page_break(doc)
    doc.add_heading("9. Running SPARTA with Codex", level=1)
    doc.add_paragraph(
        "Codex is most useful here as a guarded laboratory assistant: it can read the repository contract, run tests, inspect logs, and explain failures. "
        "It should not silently install privileged packages, alter the benchmark, or declare a learning run validated."
    )
    doc.add_heading("9.1 Install and start the CLI", level=2)
    add_code(doc, "curl -fsSL https://chatgpt.com/codex/install.sh | sh\ncd DSMC_Python/sparta_cavity_mohammadzadeh\ncodex")
    doc.add_paragraph(
        "In the interactive session, ask Codex to read `AGENTS.md` and `VALIDATION_STATUS.md` first. Then request the smoke workflow and review each command before extending it."
    )
    doc.add_heading("9.2 Guarded non-interactive run", level=2)
    add_code(doc, "# Build SPARTA once\nbash scripts/install_sparta_linux.sh serial\n\n# Let Codex run the bounded verification prompt\nbash scripts/run_with_codex.sh\n\n# Read the saved summary\nless runs/codex_smoke_report.md")
    doc.add_paragraph(
        "The wrapper invokes `codex exec --sandbox workspace-write`. The bundled prompt runs unit tests and smoke only, stops if SPARTA is absent, forbids `sudo`, and saves a concise PASS/FAIL report."
    )
    doc.add_heading("9.3 A good interactive prompt", level=2)
    add_code(doc, "Read AGENTS.md, README.md, and VALIDATION_STATUS.md. Run the unit tests and\nthe serial smoke case. Inspect the SPARTA log, generated metadata, and validation\nmetrics. Do not change constants, use sudo, or start production. Explain why this\nrun is or is not technically healthy, and give exactly one next command.")
    doc.add_heading("9.4 What Codex may change", level=2)
    add_table(
        doc,
        ["Allowed after review", "Requires explicit scientific decision"],
        [
            ["Fix path handling, parser errors, tests, documentation, and Slurm syntax.", "Gas constants, wall model, Kn definition, reference data, filters, validation gates."],
            ["Add a new, separately named resolution or seed matrix.", "Replace or relabel the paper baseline."],
        ],
        [4680, 4680],
    )

    page_break(doc)
    doc.add_heading("10. Student exercises", level=1)
    exercises = [
        ("Audit the deck", "For the student preset, recompute λ, n, fnum, Δx/λ, and Δt/τ from `case_metadata.json`. Explain every discrepancy from hand rounding."),
        ("Resolution ladder", "Run smoke and student. Plot both against the PRE data and explain why faster is not equivalent to more credible."),
        ("Repeat seeds", "Add three student seeds without overwriting outputs. Compute mean and 95% confidence intervals at the PRE x/L coordinates."),
        ("Time-step sensitivity", "Repeat with Δt/2 at fixed grid and population. Compare slip and temperature RMSE."),
        ("VHS versus VSS", "Create a sensitivity-only α = 1.4 data file. Hold numerical settings fixed and discuss physical—not just visual—differences."),
        ("Knudsen sweep", "Generate Kn = 0.05 and 0.005 cases. Choose grids based on Δx/λ rather than copying Nx blindly."),
        ("Raw versus filtered", "Implement the paper's five-neighbor filter as a new output column. Preserve and plot raw data beside it."),
        ("Parallel efficiency", "Run one identical student problem on 1, 2, 4, and 8 MPI tasks. Report speedup, efficiency, and particle imbalance."),
        ("Independent solver", "Compare SPARTA with the repository Python NTC-PreScan implementation using the same gas and wall assumptions."),
        ("Codex audit", "Ask Codex to identify every statement that would be needed before writing 'validated'. Manually verify its evidence references."),
    ]
    for index, (title_text, body) in enumerate(exercises, 1):
        if index == 7:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        run = p.add_run(f"{index}. {title_text}")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)
        p2 = doc.add_paragraph(body)
        p2.paragraph_format.left_indent = Inches(0.25)
    add_callout(
        doc,
        "Deliverable standard",
        "Every exercise submission must include its exact command, metadata JSON, relevant log excerpt, plot, numerical metric, and one paragraph separating numerical evidence from physical interpretation.",
    )

    page_break(doc)
    doc.add_heading("11. Troubleshooting", level=1)
    add_table(
        doc,
        ["Symptom", "Likely cause", "Action"],
        [
            ["SPARTA binary not found", "Local build was not run or mode differs", "Run install script for serial or mpi; check SPARTA_BIN."],
            ["`mpirun` missing", "MPI packages/modules absent", "Install OpenMPI locally or load the site MPI module."],
            ["Zero/near-zero collisions", "Density, fnum, species, or collision model wrong", "Audit metadata and `argon.vss`; inspect log counts."],
            ["No lid motion", "Moving wall assigned to wrong face", "Confirm `bound_modify yhi collide lid`."],
            ["Post-processor reports missing column", "Dump command changed", "Restore `id xc yc f_flowavg[*]` or update parser and tests together."],
            ["Temperature is noisy", "Too few particles/samples; transient contamination", "Increase population/warm-up/window; repeat seeds."],
            ["Slurm rejects GPU request", "CPU partition with GPU GRES or CPU-only binary", "Use supplied CPU/MPI template; build Kokkos/CUDA separately."],
            ["Codex asks for sudo", "System prerequisite is absent", "Stop automation; install through approved system/cluster process."],
        ],
        [2300, 3000, 4060],
    )
    doc.add_heading("11.1 Minimal health checklist", level=2)
    for item in (
        "Unit test is green and generated metadata carries the correct evidence label.",
        "SPARTA log reports nonzero boundary collisions and gas collisions.",
        "The final grid dump contains the expected averaged fields.",
        "The profile has one row per top grid cell and spans the domain monotonically.",
        "A failed learning-resolution metric is reported, not suppressed.",
    ):
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("References and reproducibility resources", level=1)
    refs = [
        ("Mohammadzadeh, A., Roohi, E., Niazmand, H., Stefanov, S., & Myong, R. S. (2012). Thermal and second-law analysis of a micro- or nanocavity using direct-simulation Monte Carlo. Physical Review E, 85, 056310. DOI: 10.1103/PhysRevE.85.056310.", "https://doi.org/10.1103/PhysRevE.85.056310"),
        ("SPARTA official documentation: Getting Started / build and run commands.", "https://sparta.github.io/doc/Section_start.html"),
        ("SPARTA boundary and bound_modify documentation.", "https://sparta.github.io/doc/boundary.html"),
        ("SPARTA diffuse and translating surf_collide documentation.", "https://sparta.github.io/doc/surf_collide.html"),
        ("SPARTA gas-collision documentation.", "https://sparta.github.io/doc/collide.html"),
        ("SPARTA official source repository.", "https://github.com/sparta/sparta"),
        ("OpenAI Codex CLI documentation.", "https://developers.openai.com/codex/cli/"),
        ("OpenAI Codex CLI command reference.", "https://developers.openai.com/codex/cli/reference/"),
        ("Companion book repository and executable case.", "https://github.com/Ehsan-Roohi/DSMC_Python/tree/agent/validated-dsmc-cavity/sparta_cavity_mohammadzadeh"),
    ]
    for idx, (text, url) in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(f"[{idx}] {text} ")
        add_hyperlink(p, url, url)
    doc.add_heading("Reproducibility record for this draft", level=2)
    add_table(
        doc,
        ["Item", "Recorded value"],
        [
            ["SPARTA upstream commit", "912c9e163c38ea5c3562d039e65215f6e2a4f3f8"],
            ["SPARTA executable banner", "24 Sep 2025"],
            ["Executed teaching preset", "student, 50 × 50, 12 particles/cell, seed 20260803"],
            ["Pilot metrics", f"slip RMSE {metrics['slip_rmse']:.4f}; temperature RMSE {metrics['temperature_rmse_K']:.2f} K"],
            ["Scientific status", "Workflow verified; production validation pending"],
        ],
        [2900, 6460],
    )
    add_callout(
        doc,
        "Recommended next manuscript step",
        "After the Unity production matrix finishes, replace Figure 2 with repeated-seed production means and confidence intervals. "
        "Keep the present pilot as an instructor note or online exercise so readers can see why evidence labels matter.",
        PALE_BLUE,
    )

    format_inline_code(doc)
    DOCS.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
