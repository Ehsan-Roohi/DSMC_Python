#!/usr/bin/env python3
"""Build the visually verified Word report for the SPARTA cavity tutorial."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results" / "production_62778322"
REPORT_ASSETS = Path(__file__).resolve().parent / "assets"
OUTPUT = Path(__file__).resolve().parent / "SPARTA_Lid_Driven_Cavity_Tutorial.docx"
GITHUB_URL = (
    "https://github.com/Ehsan-Roohi/DSMC_Python/tree/"
    "agent/sparta-lid-driven-cavity-tutorial/sparta_lid_driven_cavity_tutorial"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17324D"
MUTED = "5F6B76"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "000000"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    if sum(widths) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA, got {sum(widths)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.append(color)
    run_pr.append(underline)
    run.append(run_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def add_code_block(doc: Document, lines: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.10)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    run = paragraph.add_run(lines.strip("\n"))
    set_run_font(run, name="Consolas", size=8.5, color=BLACK)


def add_paragraph(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        run = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0]
    repeat_header(header)
    for index, label in enumerate(headers):
        set_cell_shading(header.cells[index], BLUE)
        paragraph = header.cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(label)
        set_run_font(run, size=9, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            if row_index % 2:
                set_cell_shading(cells[column_index], "F8FAFC")
            paragraph = cells[column_index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            run = paragraph.add_run(value)
            set_run_font(run, size=9)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "SPARTA Tutorial  |  Rarefied Lid-Driven Cavity"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def build() -> Path:
    summary = json.loads((RESULTS / "ensemble_summary.json").read_text(encoding="utf-8"))
    doc = Document()
    configure_document(doc)

    opening = doc.add_paragraph()
    opening.alignment = WD_ALIGN_PARAGRAPH.CENTER
    opening.paragraph_format.space_before = Pt(52)
    opening.paragraph_format.space_after = Pt(8)
    run = opening.add_run("RUNNING A RAREFIED\nLID-DRIVEN CAVITY WITH SPARTA")
    set_run_font(run, size=25, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run("A reproducible DSMC tutorial with serial, MPI, and Unity workflows")
    set_run_font(run, size=13.5, color=DARK_BLUE, italic=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(18)
    run = meta.add_run("Teaching report  |  11 August 2026  |  Kn = 0.1")
    set_run_font(run, size=10, color=MUTED, bold=True)

    link = doc.add_paragraph()
    link.alignment = WD_ALIGN_PARAGRAPH.CENTER
    link.paragraph_format.space_after = Pt(26)
    add_hyperlink(link, "Open the complete code and data on GitHub", GITHUB_URL)

    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "Outcome: generate a dimensional SPARTA input, run a fully diffuse VHS-argon cavity, "
        "post-process raw and display-smoothed fields, and combine independent random seeds."
    )
    set_run_font(r, size=10.5, color=INK, bold=True)
    set_table_geometry(callout, [9360])

    doc.add_heading("1. Purpose and learning objectives", level=1)
    add_paragraph(
        doc,
        "This chapter presents a complete SPARTA workflow for a two-dimensional rarefied lid-driven cavity. "
        "The emphasis is reproducibility: every dimensional input is generated by a script, the executable is built from a pinned SPARTA commit, and the post-processor reports the solver output directly."
    )
    for item in (
        "Identify the SPARTA commands that define the domain, molecular model, moving wall, and sampling operation.",
        "Run smoke, classroom, and production presets without editing dimensional numbers by hand.",
        "Read SPARTA timing, particle, and collision diagnostics from log.cavity.",
        "Combine three independent seeds and distinguish raw DSMC output from labelled display smoothing.",
    ):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)

    section_two = doc.add_heading("2. Physical and numerical setup", level=1)
    section_two.paragraph_format.page_break_before = True
    add_paragraph(
        doc,
        "The domain is a square cavity of side L = 1.0 × 10⁻⁶ m. The stationary walls and translating lid are fully diffuse at 300 K, and the lid moves at 100 m/s in the positive x direction. The gas is monatomic argon."
    )
    config_rows = [
        ["Knudsen number", "0.1", "Mean free path λ = 1.0 × 10⁻⁷ m"],
        ["Grid", "200 × 200", "Δx/λ = 0.05"],
        ["Simulator population", "32 particles/cell", "1,280,000 nominal particles"],
        ["Argon mass", "6.63 × 10⁻²⁶ kg", "Monatomic species"],
        ["VHS parameters", "dref = 4.17 × 10⁻¹⁰ m", "ω = 0.81; Tref = 273 K; α = 1"],
        ["Time step", "8.2568 × 10⁻¹³ s", "Δt/τc ≈ 0.00292"],
        ["Averaging", "14,000 + 26,000 steps", "Stride = 10 steps"],
    ]
    add_table(doc, ["Quantity", "Production value", "Interpretation"], config_rows, [2600, 2600, 4160])
    add_paragraph(
        doc,
        "SPARTA exposes the VHS limit through its vss collision style. The collision-data file sets α = 1.0; this produces isotropic hard-sphere angular scattering while retaining the variable collision diameter controlled by ω."
    )

    doc.add_heading("3. The generated SPARTA input", level=1)
    doc.add_heading("3.1 Domain and grid", level=2)
    add_code_block(
        doc,
        """dimension            2
boundary             s s p
create_box           0.0 1.0e-6 0.0 1.0e-6 -0.5 0.5
create_grid          200 200 1""",
    )
    doc.add_heading("3.2 Gas, walls, and molecular collisions", level=2)
    add_code_block(
        doc,
        """species              argon.species Ar
mixture              gas Ar nrho 1.294383653e25 temp 300
surf_collide         fixed diffuse 300 1.0
surf_collide         lid diffuse 300 1.0 translate 100 0.0 0.0
bound_modify         xlo xhi ylo collide fixed
bound_modify         yhi collide lid
collide              vss gas argon.vss""",
    )
    doc.add_heading("3.3 Warm-up and sampling", level=2)
    add_code_block(
        doc,
        """run                  14000
reset_timestep       0
compute              flow grid all gas nrho u v w temp
fix                  flowavg ave/grid all 10 1 10 c_flow[*] ave running
dump                 fields grid all 26000 grid.final.* id xc yc f_flowavg[*]
run                  26000""",
    )
    add_paragraph(
        doc,
        "The first block develops the flow before statistics are accumulated. The second block samples number density, velocity components, and translational temperature. The final dump therefore contains running time averages rather than an instantaneous particle snapshot."
    )

    doc.add_heading("4. Build, test, and run", level=1)
    doc.add_heading("4.1 Download and build", level=2)
    add_code_block(
        doc,
        """git clone --branch agent/sparta-lid-driven-cavity-tutorial --single-branch \\
  https://github.com/Ehsan-Roohi/DSMC_Python.git
cd DSMC_Python/sparta_lid_driven_cavity_tutorial
bash scripts/install_sparta_linux.sh serial""",
    )
    doc.add_heading("4.2 Unit tests and smoke run", level=2)
    add_code_block(
        doc,
        """python3 -m unittest discover -s tests -v
bash scripts/run_case.sh smoke serial
bash scripts/run_case.sh tutorial serial""",
    )
    doc.add_heading("4.3 MPI production run", level=2)
    add_code_block(
        doc,
        "SEED=20260803 MPI_RANKS=16 bash scripts/run_case.sh production mpi",
    )
    add_paragraph(
        doc,
        "The runner refuses to overwrite an existing run directory. Each run keeps its generated input, model data, metadata, log, grid dump, profiles, figures, and JSON summary together."
    )

    doc.add_heading("5. Unity/Slurm workflow", level=1)
    add_paragraph(
        doc,
        "The Unity helper submits a short build job followed by a three-task production array. Each array task uses one seed and 16 MPI CPU ranks."
    )
    add_code_block(
        doc,
        """bash hpc/submit_unity.sh
source LAST_SPARTA_TUTORIAL_JOBS.env
squeue -j "${BUILD_JOB_ID},${ARRAY_JOB_ID}"
sacct -X -j "${BUILD_JOB_ID},${ARRAY_JOB_ID}" \\
  --format=JobID%24,JobName%22,State,ExitCode,Elapsed,MaxRSS,NodeList%20""",
    )
    add_paragraph(
        doc,
        "The array script performs an MPI preflight, feeds the input deck on standard input to avoid launcher-option ambiguity, checks for grid.final.00026000, and then runs the standalone post-processor."
    )

    doc.add_heading("6. Post-processing and ensemble statistics", level=1)
    add_code_block(
        doc,
        """python3 scripts/ensemble_postprocess.py \\
  runs/production_kn01_seed_20260803 \\
  runs/production_kn01_seed_20260819 \\
  runs/production_kn01_seed_20260831 \\
  --output runs/production_kn01_ensemble""",
    )
    add_paragraph(
        doc,
        "The ensemble script calculates cellwise means and two-sided 95% Student-t intervals across independent seeds. Raw lid values remain in CSV form. An 11-cell moving average is plotted as a labelled readability aid; field panels use a sigma-one-cell Gaussian display filter."
    )

    doc.add_heading("7. Completed SPARTA results", level=1)
    add_paragraph(
        doc,
        "Slurm array 62778322 completed all three production seeds on 16 MPI CPU ranks. Every seed produced the expected final dump at sampling step 26,000, and no particle became stuck."
    )
    runtime_rows = []
    for seed in summary["seeds"]:
        values = summary["runtime_by_seed"][str(seed)]
        runtime_rows.append(
            [
                str(seed),
                f"{values['warmup_loop_seconds']:.3f}",
                f"{values['sampling_loop_seconds']:.3f}",
                f"{values['sampling_timesteps_per_second']:.3f}",
                f"{values['sampling_collisions']:,}",
                str(values["particles_stuck"]),
            ]
        )
    add_table(
        doc,
        ["Seed", "Warm-up (s)", "Sampling (s)", "Steps/s", "Collisions", "Stuck"],
        runtime_rows,
        [1260, 1500, 1500, 1400, 2400, 1300],
    )
    result_rows = [
        ["Maximum ensemble-mean speed", f"{summary['domain_mean_speed_max_m_per_s']:.3f} m/s"],
        [
            "Ensemble-mean domain temperature",
            f"{summary['domain_temperature_mean_min_K']:.3f}–{summary['domain_temperature_mean_max_K']:.3f} K",
        ],
        ["Mean lid-adjacent temperature", f"{summary['lid_temperature_mean_K']:.3f} K"],
        [
            "11-cell lid temperature range",
            f"{summary['lid_temperature_11cell_min_K']:.3f}–{summary['lid_temperature_11cell_max_K']:.3f} K",
        ],
        [
            "Mean lid cellwise seed standard deviation",
            f"{summary['mean_lid_seed_standard_deviation_temperature_K']:.3f} K",
        ],
    ]
    add_table(doc, ["Reported quantity", "Three-seed result"], result_rows, [5700, 3660])

    figure = doc.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.keep_with_next = True
    figure.add_run().add_picture(
        str(REPORT_ASSETS / "ensemble_lid_profiles_report.jpg"), width=Inches(6.45)
    )
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = False
    caption.paragraph_format.space_after = Pt(8)
    run = caption.add_run(
        "Figure 1. Three-seed lid profiles. Thin gray: raw ensemble mean; shaded region: 95% seed interval; heavy curve: labelled 11-cell display average."
    )
    set_run_font(run, size=9, color=MUTED, italic=True)

    figure = doc.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.keep_with_next = True
    figure.add_run().add_picture(
        str(REPORT_ASSETS / "ensemble_fields_report.jpg"), width=Inches(6.45)
    )
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    run = caption.add_run(
        "Figure 2. Ensemble-mean temperature (left) and speed (right), both with continuous streamlines. A sigma-one-cell filter is used only for display."
    )
    set_run_font(run, size=9, color=MUTED, italic=True)

    add_paragraph(
        doc,
        "The translating lid creates one dominant clockwise recirculating vortex. The maximum cell-centered speed is below the imposed lid speed because the sampled gas lies inside the cavity. Despite isothermal walls, the gas temperature is nonuniform: shear work, molecular energy transport, local expansion, and nonequilibrium effects redistribute translational energy."
    )
    add_paragraph(
        doc,
        "The variation in measured wall-clock performance among seeds reflects different allocated nodes and runtime conditions. The collision counts and macroscopic patterns are reproducible, so timing variability should be interpreted as a computing-performance observation rather than a change in the physical case."
    )

    doc.add_heading("8. Reading the output files", level=1)
    output_rows = [
        ["case_metadata.json", "Complete dimensional and numerical setup"],
        ["log.cavity", "Particle count, collisions, loop time, and performance"],
        ["grid.final.00026000", "Time-averaged cell fields at the end of sampling"],
        ["lid_profile_raw.csv", "Unsmoothed top-cell-center profile"],
        ["lid_profile_11cell.csv", "Labelled display average"],
        ["run_summary.json", "Seed-specific solver and field statistics"],
        ["ensemble_lid_profile.csv", "Three-seed mean and 95% interval"],
    ]
    add_table(doc, ["File", "Use"], output_rows, [3200, 6160])

    doc.add_heading("9. Suggested classroom exercise", level=1)
    exercises = [
        "Run the unit tests and smoke case, then identify which errors each stage can catch.",
        "Inspect argon.vss and explain why alpha = 1.0 makes the collision law VHS.",
        "Compare metadata from the smoke, tutorial, and production presets.",
        "Read log.cavity and locate particle count, attempted collisions, completed collisions, loop time, and MPI performance.",
        "Plot the raw lid profile before adding the labelled display curve.",
        "Repeat with a new random seed and quantify stochastic variability with the ensemble script.",
    ]
    for item in exercises:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        set_run_font(run)

    doc.add_heading("10. Reproducibility link", level=1)
    add_paragraph(
        doc,
        "The complete input generator, molecular data, local and Unity runners, tests, post-processing scripts, compact results, and this report are maintained together at:"
    )
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(10)
    add_hyperlink(paragraph, GITHUB_URL, GITHUB_URL)

    doc.core_properties.title = "Running a Rarefied Lid-Driven Cavity with SPARTA"
    doc.core_properties.subject = "SPARTA DSMC teaching tutorial"
    doc.core_properties.author = "Ehsan Roohi"
    doc.core_properties.keywords = "SPARTA, DSMC, lid-driven cavity, VHS, MPI, tutorial"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
